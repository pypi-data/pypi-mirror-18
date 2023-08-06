#!/usr/bin/python

import click
import subprocess
import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class AutomatedSearch(object):
    CMD_BASE = 'webkit2png -F -o temp --delay {delay} --selector div#search "https://www.google.co.uk/#q={query}"'
    TMP_NAME = "temp-full.png"

    def __init__(self, query, name, num_pages, doc_margin, delay):
        self.query = query
        self.name = name
        self.num_pages = num_pages
        self.doc_margin = doc_margin
        self.delay = delay
        if not os.path.exists(self.name):
            os.makedirs(self.name)

    def full_query(self, i):
        return '{}&start={}0'.format('+'.join(self.query.split()), i)

    def full_png_name(self, i):
        return "{0}/{0}-{1}.png".format(self.name, i+1)

    def full_docx_name(self):
        return "{0}/{0}.docx".format(self.name)

    def command(self, i):
        return AutomatedSearch.CMD_BASE.format(query=self.full_query(i), delay=self.delay)

    def take_png_shots(self):
        i = 0
        retry = 3
        while i < self.num_pages:
            try:
                self.take_png_shot(i)
            except Exception as e:
                retry -= 1
                if retry == 0:
                    raise
                print "Failed to generate shot #{}, retrying...".format(i + 1)
            else:
                i += 1
                retry = 3

    def take_png_shot(self, i):
        cmd = self.command(i)
        print "Running command '{}'".format(cmd)
        subprocess.Popen(cmd, shell=True).wait()
        os.rename(AutomatedSearch.TMP_NAME, self.full_png_name(i))

    def make_docx(self):
        self.take_png_shots()
        document = Document()
        section = document.sections[-1]
        section.top_margin = Cm(self.doc_margin)
        section.bottom_margin = Cm(self.doc_margin)
        for i in range(self.num_pages):
            document.add_picture(
                self.full_png_name(i),
                height=section.page_height - section.top_margin - section.bottom_margin)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.save(self.full_docx_name())


@click.command()
@click.option('--query', default="test", prompt='Search query', help="Query to send to google")
@click.option('--name', default='search', prompt='Name of search', help="name used for output files")
@click.option('--num_pages', default=3, help="number of pages to retrieve")
@click.option('--doc_margin', default=1, help="margin in cm in docx")
@click.option('--delay', default=1, help="delay applied to query, increase if the connection is bad")
def main(*args, **kwargs):
    auto_search = AutomatedSearch(*args, **kwargs)
    auto_search.make_docx()


if __name__ == "__main__":
    main()
